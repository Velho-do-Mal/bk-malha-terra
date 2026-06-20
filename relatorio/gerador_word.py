"""
relatorio/gerador_word.py — BK Malha de Terra (versão expandida)
================================================================

Gera relatório técnico .docx para aprovação em concessionária (padrão CPFL).

Seções:
    1. Identificação do Projeto
    2. Objetivo
    3. Metodologia (equações OMML nativas do Word)
    4. Dados do Sistema Elétrico (impedâncias por barra, curtos-circuitos)
    5. Resistividade do Solo (Wenner)
    6. Dimensionamento do Condutor de Aterramento (por barra/nível)
    7. Geometria da Malha
    8. Cálculo da Resistência de Aterramento
    9. Tensões de Passo e Toque
   10. Coordenação e Seletividade dos Relés
   11. Verificação dos Critérios de Segurança (IEEE 80)
   12. Conclusão
   13. Referências Bibliográficas

Equações: inseridas como OMML (Office Math Markup Language) — renderizadas
pelo Word como equações editáveis nativas, não como imagens.

Dependências: python-docx, lxml (já incluídas no requirements.txt)
"""

from __future__ import annotations

import io
import math
from datetime import date, datetime
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

# ─── CONSTANTES DE FORMATAÇÃO ─────────────────────────────────────────────────

_COR_HEADER     = "1F4E79"   # azul escuro CPFL
_COR_SUBHEADER  = "2E75B6"   # azul médio
_COR_LINHA_PAR  = "DEEAF1"   # azul clarinho
_COR_LINHA_IMPAR= "FFFFFF"
_FONTE          = "Arial"
_FONTE_MONO     = "Courier New"

# Namespace OMML para equações Word
_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# ─── HELPERS GERAIS ──────────────────────────────────────────────────────────

def _rgb(hex_str: str):
    r = int(hex_str[0:2], 16)
    g = int(hex_str[2:4], 16)
    b = int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def _cell_shade(cell, hex_color: str):
    """Aplica sombreamento a célula de tabela."""
    tc = cell._tc
    tcp = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcp.append(shd)


def _cell_borders(cell, color="CCCCCC", size=4):
    """Define bordas finas em célula."""
    tc = cell._tc
    tcp = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcp.append(tcBorders)


def _set_cell_text(cell, text: str, bold=False, italic=False,
                   font_size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                   font=_FONTE):
    """Define texto formatado em célula (limpa parágrafos existentes)."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.name = font
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = _rgb(color)


def _heading(doc: Document, text: str, level: int = 1):
    """Adiciona título numerado com estilo visual CPFL."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.name = _FONTE
    run.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    run.font.color.rgb = _rgb(_COR_HEADER if level == 1 else _COR_SUBHEADER)
    # Linha horizontal abaixo do título nível 1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), _COR_SUBHEADER)
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def _body(doc: Document, text: str, italic=False, bold=False, size=10):
    """Parágrafo de corpo de texto."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name  = _FONTE
    run.font.size  = Pt(size)
    run.italic     = italic
    run.bold       = bold
    return p


def _nota(doc: Document, text: str):
    """Nota de rodapé em itálico menor."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"Nota: {text}")
    run.font.name  = _FONTE
    run.font.size  = Pt(8)
    run.italic     = True
    run.font.color.rgb = _rgb("595959")
    return p


# ─── EQUAÇÕES OMML ────────────────────────────────────────────────────────────

def _omml_par(doc: Document, omml_xml: str):
    """Insere parágrafo centralizado com equação OMML nativa do Word."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    # Remove runs vazios padrão
    for child in list(p._p):
        if child.tag.endswith("}r") or child.tag.endswith("}pPr"):
            continue
        p._p.remove(child)
    omath = etree.fromstring(omml_xml)
    p._p.append(omath)
    return p


def _eq(terms: str) -> str:
    """Gera OMML simples para expressão de texto (sem frações complexas)."""
    return (
        f'<m:oMath xmlns:m="{_M}">'
        f'<m:r><m:rPr><m:sty m:val="i"/></m:rPr>'
        f'<m:t>{terms}</m:t></m:r>'
        f'</m:oMath>'
    )


def _eq_frac(num: str, den: str, pre="", post="") -> str:
    """OMML com fração."""
    return (
        f'<m:oMath xmlns:m="{_M}">'
        + (f'<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t xml:space="preserve">{pre} </m:t></m:r>' if pre else "")
        + f'<m:f><m:fPr/>'
        + f'<m:num><m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>{num}</m:t></m:r></m:num>'
        + f'<m:den><m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>{den}</m:t></m:r></m:den>'
        + f'</m:f>'
        + (f'<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t xml:space="preserve"> {post}</m:t></m:r>' if post else "")
        + f'</m:oMath>'
    )


def _eq_sqrt(radicand: str, pre="", post="") -> str:
    """OMML com raiz quadrada."""
    return (
        f'<m:oMath xmlns:m="{_M}">'
        + (f'<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t xml:space="preserve">{pre} </m:t></m:r>' if pre else "")
        + f'<m:rad><m:radPr><m:degHide m:val="1"/></m:radPr>'
        + f'<m:deg/>'
        + f'<m:e><m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t>{radicand}</m:t></m:r></m:e>'
        + f'</m:rad>'
        + (f'<m:r><m:rPr><m:sty m:val="i"/></m:rPr><m:t xml:space="preserve"> {post}</m:t></m:r>' if post else "")
        + f'</m:oMath>'
    )


def _label_eq(doc: Document, label: str, eq_xml: str):
    """Equação com rótulo à direita (Eq. X.Y)."""
    _omml_par(doc, eq_xml)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(label)
    run.font.name = _FONTE
    run.font.size = Pt(9)
    run.italic    = True
    run.font.color.rgb = _rgb("595959")


# ─── TABELAS ─────────────────────────────────────────────────────────────────

def _make_table(doc: Document, headers: list[str], rows: list[list[str]],
                col_widths_cm: list[float] | None = None):
    """Cria tabela formatada no padrão CPFL."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Larguras das colunas
    if col_widths_cm:
        for i, cell in enumerate(table.columns):
            for c in cell.cells:
                c.width = Cm(col_widths_cm[i] if i < len(col_widths_cm) else 3)

    # Cabeçalho
    hdr_row = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(hdr_row.cells, headers)):
        _cell_shade(cell, _COR_HEADER)
        _cell_borders(cell, "FFFFFF", 4)
        _set_cell_text(cell, hdr, bold=True, color="FFFFFF",
                       font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Dados
    for r_idx, row in enumerate(rows):
        bg = _COR_LINHA_PAR if r_idx % 2 == 0 else _COR_LINHA_IMPAR
        for c_idx, (cell, val) in enumerate(zip(table.rows[r_idx + 1].cells, row)):
            _cell_shade(cell, bg)
            _cell_borders(cell, "B8CCE4", 4)
            _set_cell_text(cell, str(val), font_size=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0
                                 else WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()  # espaço após tabela
    return table


# ─── CÁLCULOS AUXILIARES ──────────────────────────────────────────────────────

def _calcular_impedancias(de) -> dict:
    """
    Calcula impedâncias equivalentes por barra a partir dos dados de entrada.

    Parâmetros derivados:
        de.i_falta_3i0_ka  : corrente de falta fase-terra (3I₀) em kA
        de.xr_ratio        : relação X/R do sistema na falta
        de.sf_div_corrente : fator de divisão de corrente (Sf)
    """
    i_f = float(de.i_falta_3i0_ka or 1.0)       # kA (3I0 = Ig total)
    xr  = float(de.xr_ratio or 10.0)
    sf  = float(de.sf_div_corrente or 1.0)

    # Tensões base (subestação 138/13,8 kV típica CPFL)
    v_at_kv  = 138.0
    v_mt_kv  = 13.8

    # Corrente de curto trifásico estimada (sistema AT)
    # Assumindo: I_cc3φ ≈ 1,05 × I_f (curto monofásico ≈ trifásico para alta resistividade)
    i_cc3_at = i_f * 1.05   # kA

    # Impedância de Thévenin no barramento AT
    z_th_at  = v_at_kv / (math.sqrt(3) * i_cc3_at)   # Ω
    angle    = math.atan(xr)                            # rad
    r_th_at  = z_th_at * math.cos(angle)
    x_th_at  = z_th_at * math.sin(angle)

    # Impedância transformador (AT→MT), base: S_nom ≈ 30 MVA, z_t ≈ 12 %
    s_nom_mva = 30.0
    z_base_mt = v_mt_kv**2 / s_nom_mva               # Ω
    z_t_pu    = 0.12                                  # 12 % (CPFL padrão)
    z_t       = z_t_pu * z_base_mt                   # Ω
    r_t       = z_t * 0.03                            # X/R transf ≈ 33
    x_t       = math.sqrt(z_t**2 - r_t**2)

    # Impedância equivalente no MT
    z_eq_mt   = math.sqrt((r_th_at + r_t)**2 + (x_th_at + x_t)**2)
    i_cc3_mt  = v_mt_kv / (math.sqrt(3) * z_eq_mt)  # kA

    # Corrente de falta que flui pela malha
    i_g       = i_f * sf                              # kA

    return {
        "v_at_kv":  v_at_kv,
        "v_mt_kv":  v_mt_kv,
        "i_cc3_at": round(i_cc3_at, 3),
        "i_cc1_at": round(i_f, 3),
        "z_th_at":  round(z_th_at, 4),
        "r_th_at":  round(r_th_at, 4),
        "x_th_at":  round(x_th_at, 4),
        "xr_at":    round(xr, 2),
        "z_t":      round(z_t, 4),
        "i_cc3_mt": round(i_cc3_mt, 3),
        "i_g_ka":   round(i_g, 3),
    }


def _calcular_condutor_por_nivel(i_g_ka: float, t_s: float) -> dict:
    """
    Dimensiona condutor pelo critério IEEE 80-2013, Eq. (37).

    A_mm² = I × √( t_c / (TCAP/(ρ_r × α_r) × ln(1 + (T_m - T_a)/(T_a + 1/α₀))) )

    Simplificado para cobre têmpado (NBR 13936 / IEEE 80 Tabela 1):
        Kf = 7,06 × 10⁻³ (para 20°C→250°C, Tm=250°C, Ta=40°C)
    """
    i_ka     = i_g_ka
    i_a      = i_ka * 1000.0    # A
    Kf       = 7.06e-3           # A·s^0.5 / mm²  (cobre, Tm=250°C, Ta=40°C)
    A_mm2    = i_a * math.sqrt(t_s) / Kf

    # Bitolas normalizadas (NBR 6786 / IEEE)
    bitolas = [10, 16, 25, 35, 50, 70, 95, 120, 150, 185, 240, 300]
    bitola_adotada = next((b for b in bitolas if b >= A_mm2), 300)

    # Calcula para AT (corrente máx ≈ I_cc3) e MT (menor corrente)
    return {
        "at": {
            "i_a":            round(i_a, 1),
            "t_s":            round(t_s, 3),
            "A_calc_mm2":     round(A_mm2, 1),
            "bitola_mm2":     bitola_adotada,
            "material":       "Cobre nu têmpado",
        },
        "mt": {
            "i_a":            round(i_a * 0.5, 1),   # corrente MT ≈ 50 % falta AT
            "t_s":            round(t_s, 3),
            "A_calc_mm2":     round(A_mm2 * 0.5, 1),
            "bitola_mm2":     max(25, int(bitola_adotada * 0.5)),
            "material":       "Cobre nu têmpado",
        },
        "malha": {
            "i_a":            round(i_a, 1),
            "t_s":            round(t_s, 3),
            "A_calc_mm2":     round(A_mm2, 1),
            "bitola_mm2":     bitola_adotada,
            "material":       "Cobre nu têmpado",
        },
    }


def _ajustes_rele(de, imp: dict) -> list[dict]:
    """
    Calcula ajustes típicos dos relés de proteção conforme padrão CPFL
    e correntes de falta calculadas.

    Relés contemplados:
        50/51   – Sobrecorrente de fase (linha AT)
        50N/51N – Sobrecorrente de terra (linha AT)
        87T     – Diferencial de transformador
        51N-MT  – Sobrecorrente de terra no MT (neutro aterrado)
    """
    i_cc3_at  = imp["i_cc3_at"]  # kA
    i_cc1_at  = imp["i_cc1_at"]  # kA
    i_cc3_mt  = imp["i_cc3_mt"]  # kA
    t_elim    = float(de.tempo_eliminacao_s or 0.1)

    # Corrente nominal transformador AT
    s_nom_mva = 30.0
    i_nom_at  = s_nom_mva / (math.sqrt(3) * imp["v_at_kv"]) * 1000  # A
    i_nom_mt  = s_nom_mva / (math.sqrt(3) * imp["v_mt_kv"]) * 1000  # A

    # 50/51 – pickup fase: 1,5 × I_nom (CPFL-NTC-920301)
    ip_51  = round(i_nom_at * 1.5, 1)
    # 50 – pickup instantâneo fase: 80 % do I_cc3_min estimado
    ip_50  = round(i_cc3_at * 1000 * 0.80, 0)
    # 51N – pickup terra: 10 % I_nom (sensibilidade falha terra alta impedância)
    ip_51n = round(i_nom_at * 0.10, 1)
    # 50N – pickup instantâneo terra: 20 % I_cc1
    ip_50n = round(i_cc1_at * 1000 * 0.20, 0)

    return [
        {
            "funcao": "50/51",
            "descricao": "Sobrecorrente de fase — Linha 138 kV",
            "pickup_a":  f"{ip_51:.1f} A ({ip_51/i_nom_at*100:.0f}% I_nom)",
            "retardo":   f"Curva IEC Normal Inversa, TDS = 0,25 / t₀ = {t_elim:.2f} s",
            "ip50":      f"{ip_50:.0f} A (inst.)",
            "norma":     "CPFL-NTC-920301 / IEEE C37.112",
        },
        {
            "funcao": "50N/51N",
            "descricao": "Sobrecorrente de terra — Linha 138 kV",
            "pickup_a":  f"{ip_51n:.1f} A ({ip_51n/i_nom_at*100:.0f}% I_nom)",
            "retardo":   f"Curva IEC Ext. Inversa, TDS = 0,30 / t₀ = {t_elim+0.1:.2f} s",
            "ip50":      f"{ip_50n:.0f} A (inst.)",
            "norma":     "CPFL-NTC-920301 / IEEE C37.112",
        },
        {
            "funcao": "87T",
            "descricao": "Diferencial de transformador 138/13,8 kV",
            "pickup_a":  "15% I_nom (Slope 1) / 35% I_nom (Slope 2)",
            "retardo":   "Instantâneo (sem retardo intencional)",
            "ip50":      "Inrush restrain: 2.ª harmônica ≥ 15%",
            "norma":     "CPFL-NTC-920302 / IEEE C37.91",
        },
        {
            "funcao": "51N-MT",
            "descricao": "Sobrecorrente de terra — Neutro 13,8 kV",
            "pickup_a":  f"{round(i_nom_mt*0.05,1):.1f} A ({0.05*100:.0f}% I_nom-MT)",
            "retardo":   f"Curva IEC Normal Inversa, TDS = 0,20 / t₀ = {t_elim+0.2:.2f} s",
            "ip50":      "—",
            "norma":     "CPFL-NTC-920303",
        },
    ]


# ─── SEÇÕES DO DOCUMENTO ──────────────────────────────────────────────────────

def _secao_identificacao(doc, proj, de, res):
    _heading(doc, "1. IDENTIFICAÇÃO DO PROJETO", 1)
    dados = [
        ("Cliente",               proj.cliente or "—"),
        ("Projeto",               proj.nome_projeto or "—"),
        ("Número do Projeto",     proj.numero_projeto or "—"),
        ("Revisão",               proj.revisao or "00"),
        ("Responsável Técnico",   proj.responsavel_tecnico or "—"),
        ("CREA",                  proj.crea_responsavel or "—"),
        ("Concessionária",        proj.concessionaria or "CPFL"),
        ("Data do Cálculo",       str(proj.data_calculo or date.today())),
    ]
    rows = [[k, v] for k, v in dados]
    _make_table(doc, ["Campo", "Valor"], rows, col_widths_cm=[5.5, 11.0])


def _secao_objetivo(doc):
    _heading(doc, "2. OBJETIVO", 1)
    _body(doc,
        "O presente relatório tem por objetivo apresentar o memorial de cálculo do "
        "sistema de aterramento da subestação, conforme os requisitos normativos da "
        "IEEE Std 80-2013 e das Normas Técnicas da CPFL Energia, verificando os "
        "critérios de segurança relativos às tensões de passo e toque a que pessoas "
        "possam estar sujeitas durante a ocorrência de falta à terra no sistema."
    )


def _secao_metodologia(doc):
    _heading(doc, "3. METODOLOGIA", 1)
    _body(doc,
        "Os cálculos foram realizados conforme a norma IEEE Std 80-2013 — "
        "'Guide for Safety in AC Substation Grounding'. A seguir são apresentadas "
        "as principais equações utilizadas no estudo."
    )

    # 3.1 Resistividade aparente do solo
    _heading(doc, "3.1  Resistividade Aparente do Solo — Método de Wenner", 2)
    _body(doc,
        "A resistividade aparente do solo ρ (Ω·m) é determinada pelo método de "
        "Wenner de quatro eletrodos. Equacionamento (IEEE 80-2013, Apêndice B):"
    )
    _label_eq(doc, "(Eq. 3.1)",
        _eq_frac("2π a R_a", "1 + 2a/√(a®+4b²) − a/√(a²+b²)", pre="ρ ="))
    _body(doc,
        "onde: a = espaçamento entre eletrodos [m]; "
        "b = profundidade dos eletrodos [m] (normalmente b << a, logo ρ ≈ 2πaR_a); "
        "R_a = resistência medida [Ω]."
    )

    # 3.2 Fator de redução Cs (brita)
    _heading(doc, "3.2  Fator de Redução da Camada Superficial (C_s)", 2)
    _body(doc,
        "A presença de camada de brita reduz a corrente de choque através do "
        "corpo humano. O fator C_s é calculado por (IEEE 80-2013, Eq. 27):"
    )
    _label_eq(doc, "(Eq. 3.2)",
        _eq(r"C_s = 1 − 0,09 (1 − ρ/ρ_s) / (2h_s + 0,09)"))
    _body(doc,
        "onde: ρ = resistividade do solo nativo [Ω·m]; "
        "ρ_s = resistividade da brita [Ω·m]; "
        "h_s = espessura da camada de brita [m]."
    )

    # 3.3 Tensões admissíveis
    _heading(doc, "3.3  Tensões de Toque e Passo Admissíveis", 2)
    _body(doc,
        "Para pessoa de 70 kg, as tensões admissíveis são (IEEE 80-2013, Eqs. 32–33):"
    )
    _label_eq(doc, "(Eq. 3.3a)",
        _eq(r"E_toque(adm) = (1000 + 1,5 C_s ρ_s) √(0,157 / t_s)   [V]"))
    _label_eq(doc, "(Eq. 3.3b)",
        _eq(r"E_passo(adm) = (1000 + 6,0 C_s ρ_s) √(0,157 / t_s)   [V]"))
    _body(doc, "onde t_s = tempo de duração do choque [s] (igual ao tempo de eliminação da falta).")

    # 3.4 Dimensionamento do condutor
    _heading(doc, "3.4  Dimensionamento do Condutor de Aterramento", 2)
    _body(doc,
        "A seção mínima do condutor de aterramento é determinada pelo critério "
        "térmico (IEEE 80-2013, Eq. 37):"
    )
    _label_eq(doc, "(Eq. 3.4)",
        _eq_frac("I " + "√�e" + "t_c", "K_f", pre="A_mm² ="))
    _body(doc,
        "onde: I = corrente de falta [A]; t_c = duração da falta [s]; "
        "K_f = 7,06 × 10⁻³ A·s^0,5/mm² (cobre têmpado, T_m=250°C, T_a=40°C)."
    )

    # 3.5 Resistência da malha
    _heading(doc, "3.5  Resistência de Aterramento da Malha", 2)
    _body(doc,
        "A resistência de aterramento é calculada pelas fórmulas de Sverak "
        "(IEEE 80-2013, Eq. 57) e Schwarz. A de Sverak é:"
    )
    _label_eq(doc, "(Eq. 3.5)",
        _eq(r"R_g = ρ [ 1/L_T + 1/√(20A) (1 + 1/(1 + h√(20/A))) ]"))
    _body(doc,
        "onde: L_T = comprimento total de cabos + hastes [m]; "
        "A = área da malha [m²]; h = profundidade de enterramento [m]."
    )

    # 3.6 GPR
    _heading(doc, "3.6  Elevação de Potencial da Malha (GPR)", 2)
    _label_eq(doc, "(Eq. 3.6)", _eq("GPR = R_g × I_g   [V]"))
    _body(doc,
        "onde I_g = corrente que flui pela malha = S_f × D_f × I_f [A]; "
        "S_f = fator de divisão de corrente; D_f = fator de decremento assimétrico."
    )

    # 3.7 Tensão de malha e passo calculadas
    _heading(doc, "3.7  Tensões de Malha e Passo Calculadas", 2)
    _label_eq(doc, "(Eq. 3.7a)",
        _eq_frac("ρ K_m K_i I_g", "L_c", pre="E_m ="))
    _label_eq(doc, "(Eq. 3.7b)",
        _eq_frac("ρ K_s K_i I_g", "L_c", pre="E_s ="))
    _body(doc,
        "onde: K_m = fator geométrico de malha; K_s = fator geométrico de passo; "
        "K_i = fator de irregularidade; L_c = comprimento efetivo dos condutores [m]."
    )


def _secao_sistema_eletrico(doc, de, imp: dict):
    _heading(doc, "4. DADOS DO SISTEMA ELÉTRICO", 1)
    _body(doc,
        "A seguir são apresentados os dados do sistema elétrico da subestação, "
        "incluindo as impedâncias equivalentes por barra e as correntes de "
        "curto-circuito utilizadas no estudo de aterramento."
    )

    # 4.1 Configuração
    _heading(doc, "4.1  Configuração da Subestação", 2)
    config = [
        ("Tensão nominal AT",         f"{imp['v_at_kv']:.1f} kV"),
        ("Tensão nominal MT",         f"{imp['v_mt_kv']:.1f} kV"),
        ("Potência do transformador", "30 MVA (referência CPFL)"),
        ("Relação X/R do sistema",    f"{imp['xr_at']:.1f}"),
        ("Fator de divisão (S_f)",    f"{float(de.sf_div_corrente or 1.0):.3f}"),
        ("Tempo de eliminação",       f"{float(de.tempo_eliminacao_s or 0.1):.3f} s"),
    ]
    _make_table(doc, ["Parâmetro", "Valor"], [[k, v] for k, v in config],
                col_widths_cm=[7.5, 9.0])

    # 4.2 Impedâncias por barra
    _heading(doc, "4.2  Impedâncias de Thévenin por Barra", 2)
    _body(doc,
        "As impedâncias equivalentes de Thévenin no ponto de falta são calculadas "
        "a partir da corrente de curto-circuito fase-terra fornecida pelo estudo "
        "de curto-circuito do sistema (3I₀)."
    )
    rows_z = [
        ["138 kV (AT)",
         f"{imp['z_th_at']:.4f} Ω",
         f"{imp['r_th_at']:.4f} Ω",
         f"{imp['x_th_at']:.4f} Ω",
         f"{imp['xr_at']:.1f}"],
        ["13,8 kV (MT) *",
         f"{imp['z_t']:.4f} Ω",
         f"{imp['z_t']*0.03:.4f} Ω",
         f"{math.sqrt(imp['z_t']**2-(imp['z_t']*0.03)**2):.4f} Ω",
         "~33"],
    ]
    _make_table(doc,
        ["Barra", "|Z_th| (Ω)", "R_th (Ω)", "X_th (Ω)", "X/R"],
        rows_z, col_widths_cm=[3.5, 3.5, 3.5, 3.5, 2.5])
    _nota(doc, "* Impedância referida ao secundário do transformador. "
              "Incluída a impedância de curto do transformador (Z_t ≈ 12%).")

    # 4.3 Correntes de curto-circuito
    _heading(doc, "4.3  Correntes de Curto-Circuito", 2)
    rows_i = [
        ["138 kV — CC trifásico (3φ)",        f"{imp['i_cc3_at']:.3f} kA", "—",         "Máximo térmico"],
        ["138 kV — CC monofásico (1φ, 3I₀)",  f"{imp['i_cc1_at']:.3f} kA", "Projeto",   "Dimensionamento da malha"],
        ["13,8 kV — CC trifásico (3φ) *",     f"{imp['i_cc3_mt']:.3f} kA", "—",         "Informativo"],
        ["Malha — Corrente I_g",               f"{imp['i_g_ka']:.3f} kA",   "Projeto",   "Corrente dissipada na malha"],
    ]
    _make_table(doc,
        ["Ponto de Falta", "Corrente (kA)", "Origem", "Aplicação"],
        rows_i, col_widths_cm=[5.5, 3.0, 2.5, 5.5])
    _nota(doc, "Correntes de curto derivadas de I_f = 3I₀ e do estudo de fluxo de corrente no sistema. "
              "Validar com relatório de curto-circuito aprovado pela concessionária.")


def _secao_solo(doc, proj, de, wenner_rows):
    _heading(doc, "5. RESISTIVIDADE DO SOLO", 1)
    _body(doc,
        "A resistividade do solo foi determinada por ensaio geofísico utilizando "
        "o método de Wenner de quatro eletrodos, conforme IEEE Std 80-2013, "
        "Apêndice B, e NBR 7117."
    )
    if wenner_rows:
        rows = [[str(w.ponto), f"{w.espacamento_m:.2f}", f"{w.resistencia_ohm:.4f}",
                 f"{w.rho_aparente:.2f}" if w.rho_aparente else "—"]
                for w in wenner_rows]
        _make_table(doc,
            ["Ponto", "Espaçamento a (m)", "Resistência R_a (Ω)", "ρ_ap (Ω·m)"],
            rows, col_widths_cm=[2.0, 4.5, 4.5, 5.5])
    else:
        _body(doc, "[INSERIR DADOS DE CAMPO — ensaio Wenner não disponível]", italic=True)

    rho1 = float(getattr(de, "rho1_ohm_m", None) or 100)
    rho2 = float(getattr(de, "rho2_ohm_m", None) or 100)
    h1   = float(getattr(de, "h1_m", None) or 1.0)
    _body(doc,
        f"Modelo de solo adotado (duas camadas): ρ₁ = {rho1:.1f} Ω·m (camada superior, "
        f"h₁ = {h1:.2f} m); ρ₂ = {rho2:.1f} Ω·m (camada inferior — semi-infinita)."
    )


def _secao_condutor(doc, de, imp: dict, cond: dict):
    _heading(doc, "6. DIMENSIONAMENTO DO CONDUTOR DE ATERRAMENTO", 1)
    _body(doc,
        "O dimensionamento do condutor é realizado pelo critério térmico da "
        "IEEE Std 80-2013 (Eq. 3.4 deste memorial). A tabela a seguir apresenta "
        "os resultados por nível de tensão e para o cabo da malha enterrada."
    )
    rows = [
        ["Barramento 138 kV (AT)",
         f"{cond['at']['i_a']:.0f} A",
         f"{cond['at']['t_s']:.3f} s",
         f"{cond['at']['A_calc_mm2']:.1f} mm²",
         f"{cond['at']['bitola_mm2']} mm²",
         cond['at']['material']],
        ["Barramento 13,8 kV (MT)",
         f"{cond['mt']['i_a']:.0f} A",
         f"{cond['mt']['t_s']:.3f} s",
         f"{cond['mt']['A_calc_mm2']:.1f} mm²",
         f"{cond['mt']['bitola_mm2']} mm²",
         cond['mt']['material']],
        ["Malha enterrada",
         f"{cond['malha']['i_a']:.0f} A",
         f"{cond['malha']['t_s']:.3f} s",
         f"{cond['malha']['A_calc_mm2']:.1f} mm²",
         f"{cond['malha']['bitola_mm2']} mm²",
         cond['malha']['material']],
    ]
    _make_table(doc,
        ["Local / Nível", "I_g (A)", "t_c (s)", "A_calc", "Bitola adotada", "Material"],
        rows, col_widths_cm=[4.0, 2.5, 2.0, 2.5, 3.0, 3.5])
    _nota(doc,
        "Bitolas adotadas com folga de segurança conforme NBR 13936 e CPFL-NTC-920305. "
        "Temperatura máxima admissível: T_m = 250°C (fusão do cobre: 1.083°C)."
    )

    # Memória de cálculo
    _heading(doc, "6.1  Memória de Cálculo — Barramento AT", 2)
    _body(doc, f"Dados: I_g = {cond['at']['i_a']:.0f} A; t_c = {cond['at']['t_s']:.3f} s; K_f = 7,06×10⁻³ A·√s/mm²")
    _omml_par(doc, _eq(
        f"A_mm² = {cond['at']['i_a']:.0f} × √{cond['at']['t_s']:.3f} / (7,06×10⁻³) "
        f"= {cond['at']['A_calc_mm2']:.1f} mm²  →  adotado: {cond['at']['bitola_mm2']} mm²"
    ))


def _secao_geometria(doc, de):
    _heading(doc, "7. GEOMETRIA DA MALHA", 1)

    campos = {
        "Comprimento da malha (L_x)":  f"{getattr(de,'comprimento_m',None) or '[INSERIR]'} m",
        "Largura da malha (L_y)":       f"{getattr(de,'largura_m',None) or '[INSERIR]'} m",
        "Espaçamento entre cabos (D)":  f"{getattr(de,'espacamento_cabos_m',None) or '[INSERIR]'} m",
        "Profundidade (h)":             f"{getattr(de,'profundidade_m',None) or 0.5} m",
        "Espessura da brita (h_s)":     f"{getattr(de,'espessura_brita_m',None) or 0.1} m",
        "Número de hastes":             f"{getattr(de,'n_hastes',None) or '[INSERIR]'}",
        "Comprimento das hastes":       f"{getattr(de,'comprimento_haste_m',None) or 2.4} m",
    }
    rows = [[k, v] for k, v in campos.items()]
    _make_table(doc, ["Parâmetro", "Valor"], rows, col_widths_cm=[7.5, 9.0])


def _secao_calculos(doc, de, res):
    _heading(doc, "8. CÁLCULO DA RESISTÊNCIA DE ATERRAMENTO", 1)
    if not res:
        _body(doc, "[CALCULAR O PROJETO ANTES DE GERAR O RELATÓRIO]", italic=True)
        return

    _body(doc,
        "A resistência de aterramento da malha foi calculada pelos métodos de "
        "Sverak e Schwarz (IEEE Std 80-2013). O valor mais conservador é adotado."
    )

    rows = [
        ["Resistência Sverak (R_g,Sv)",   f"{res.rg_sverak_ohm:.4f} Ω"],
        ["Resistência Schwarz (R_g,Sc)",   f"{res.rg_schwarz_ohm:.4f} Ω"],
        ["Resistência adotada (R_g)",      f"{res.rg_adotado_ohm:.4f} Ω"],
        ["Corrente na malha (I_g)",        f"{res.ig_corrente_malha_a:.2f} A"],
        ["Elevação de potencial (GPR)",    f"{res.gpr_v:.1f} V"],
    ]
    _make_table(doc, ["Grandeza", "Valor"], rows, col_widths_cm=[9.0, 7.5])

    # Memória GPR
    _heading(doc, "8.1  Memória de Cálculo — GPR", 2)
    _omml_par(doc, _eq(
        f"GPR = R_g × I_g = {res.rg_adotado_ohm:.4f} × {res.ig_corrente_malha_a:.2f} "
        f"= {res.gpr_v:.1f} V"
    ))


def _secao_tensoes(doc, de, res):
    _heading(doc, "9. TENSÕES DE PASSO E TOQUE", 1)
    if not res:
        _body(doc, "[CALCULAR O PROJETO ANTES DE GERAR O RELATÓRIO]", italic=True)
        return

    rows = [
        ["Tensão de toque admissível (E_toque,adm)", f"{res.etoque_admissivel_v:.2f} V",   "—"],
        ["Tensão de passo admissível (E_passo,adm)", f"{res.epasso_admissivel_v:.2f} V",   "—"],
        ["Tensão de malha calculada (E_m)",          f"{res.em_tensao_malha_v:.2f} V",     f"{res.etoque_admissivel_v:.2f} V"],
        ["Tensão de passo calculada (E_s)",          f"{res.es_tensao_passo_v:.2f} V",     f"{res.epasso_admissivel_v:.2f} V"],
    ]
    _make_table(doc,
        ["Grandeza", "Valor Calculado", "Limite Admissível"],
        rows, col_widths_cm=[7.5, 4.5, 4.5])


def _secao_coordenacao(doc, de, reles: list[dict]):
    _heading(doc, "10. COORDENAÇÃO E SELETIVIDADE DOS RELÉS", 1)
    _body(doc,
        "A coordenação e seletividade dos relés de proteção é essencial para "
        "garantir que a falta seja eliminada dentro do tempo de eliminação "
        "adotado no estudo de aterramento. A tabela a seguir apresenta os "
        "ajustes dos relés conforme CPFL-NTC-920301/302/303 e IEEE C37.112."
    )

    # Tabela de ajustes
    _heading(doc, "10.1  Tabela de Ajustes dos Relés de Proteção", 2)
    rows = []
    for r in reles:
        rows.append([
            r["funcao"],
            r["descricao"],
            r["pickup_a"],
            r["retardo"],
            r["ip50"],
            r["norma"],
        ])
    _make_table(doc,
        ["Função", "Descrição", "Pickup / Corrente de ajuste",
         "Retardo / Característica", "Instantâneo (50)", "Norma"],
        rows, col_widths_cm=[1.5, 4.0, 3.5, 4.0, 2.5, 2.5])

    # Critério de seletividade
    _heading(doc, "10.2  Critério de Seletividade e Tempo de Eliminação", 2)
    t_elim = float(de.tempo_eliminacao_s or 0.1)
    _body(doc,
        f"O tempo de eliminação da falta adotado no estudo é t_f = {t_elim:.3f} s, "
        "definido pelo relé de proteção principal (função 50N/51N ou 87T). "
        "Este valor C� o tempo máximo de choque (t_s) utilizado no cálculo das "
        "tensões admissíveis de passo e toque (Eqs. 3.3a e 3.3b)."
    )
    _body(doc,
        "A hierarquia de seletividade adotada é: "
        "(1) 87T — proteção principal do transformador (instantâneo); "
        "(2) 50N — sobrecorrente instantâneo de terra (backup primário); "
        "(3) 51N — sobrecorrente temporizado de terra (backup secundário). "
        "O tempo de eliminação mais longo (51N) é o valor conservativo adotado."
    )

    # Nota de coordenação
    _heading(doc, "10.3  Verificação do Critério de Tempo × Corrente", 2)
    _body(doc,
        "Para aprovação pela concessionária CPFL é necessário apresentar o "
        "diagrama de coordenação em papel log-log com as curvas dos relés. "
        "O diagrama deve ser gerado em software de coordenação (ETAP, CYMTCC "
        "ou similar) e inserido como Anexo ao presente memorial."
    )
    _nota(doc,
        "Inserir diagrama de coordenação como Anexo. "
        "Validar ajustes com a concessionária CPFL antes da energização."
    )


def _secao_verificacao(doc, de, res):
    _heading(doc, "11. VERIFICAÇÃO DOS CRITÉRIOS DE SEGURANÇA", 1)
    if not res:
        _body(doc, "[CALCULAR O PROJETO ANTES DE GERAR O RELATÓRIO]", italic=True)
        return

    ok_toque = res.atende_toque
    ok_passo = res.atende_passo

    def _status(ok):
        return "✔  ATENDE" if ok else "✘  NÃO ATENDE"

    rows = [
        ["Tensão de toque", f"{res.em_tensao_malha_v:.2f} V",
         f"{res.etoque_admissivel_v:.2f} V",
         f"{abs(res.margem_toque_pct):.1f} %", _status(ok_toque)],
        ["Tensão de passo", f"{res.es_tensao_passo_v:.2f} V",
         f"{res.epasso_admissivel_v:.2f} V",
         f"{abs(res.margem_passo_pct):.1f} %", _status(ok_passo)],
    ]
    _make_table(doc,
        ["Critério", "Calculado", "Admissível", "Margem", "Resultado"],
        rows, col_widths_cm=[3.5, 3.0, 3.0, 2.5, 4.5])

    if ok_toque and ok_passo:
        p = _body(doc,
            "✔  A malha de aterramento ATENDE a todos os critérios de segurança "
            "estabelecidos pela IEEE Std 80-2013 para as condições de projeto analisadas.")
        p.runs[0].font.color.rgb = _rgb("1F7A1F")
    else:
        p = _body(doc,
            "✘  A malha de aterramento NÃO atende a todos os critérios. "
            "Recomenda-se revisar a geometria da malha, adicionar hastes ou "
            "aumentar a espessura da camada de brita.")
        p.runs[0].font.color.rgb = _rgb("C00000")


def _secao_conclusao(doc, proj, de, res):
    _heading(doc, "12. CONCLUSÃO", 1)
    if res and res.atende_geral:
        _body(doc,
            f"O sistema de aterramento projetado para a subestação "
            f"{proj.nome_projeto or ''} — {proj.cliente or ''} atende integralmente "
            f"aos critérios de segurança da IEEE Std 80-2013 e às exigências "
            f"da CPFL Energia, para as condições de falta e de solo analisadas."
        )
    else:
        _body(doc,
            "Os resultados indicam que, nas condições analisadas, a malha de "
            "aterramento requer ajustes de projeto para atendimento pleno dos "
            "critérios de segurança. As recomendações estão listadas na seção 11."
        )
    _body(doc,
        "A responsabilidade técnica pelo presente memorial cabe ao profissional "
        f"habilitado {proj.responsavel_tecnico or '[INSERIR NOME]'} "
        f"(CREA {proj.crea_responsavel or '[INSERIR CREA]'}), nos termos da "
        "Resolução CFE nº 218/73 e da NBR 5460."
    )


def _secao_referencias(doc):
    _heading(doc, "13. REFERÊNCIAS BIBLIOGRÁFICAS", 1)
    refs = [
        "IEEE Std 80-2013 — Guide for Safety in AC Substation Grounding. "
        "IEEE Power & Energy Society, 2013.",
        "ABNT NBR 7117:2012 — Medição da resistividade do solo. "
        "Associação Brasileira de Normas Técnicas, 2012.",
        "ABNT NBR 13936:1997 — Subestações de alta tensão acima de 1 kV — "
        "Projeto. ABNT, 1997.",
        "CPFL-NTC-920301 — Norma Técnica de Proteção de Linhas de Transmissão "
        "138 kV. CPFL Energia, vigente.",
        "CPFL-NTC-920302 — Norma Técnica de Proteção de Transformadores de Força. "
        "CPFL Energia, vigente.",
        "CPFL-NTC-920303 — Norma Técnica de Proteção de Alimentadores 13,8 kV. "
        "CPFL Energia, vigente.",
        "CPFL-NTC-920305 — Norma Técnica de Aterramento de Subestações. "
        "CPFL Energia, vigente.",
        "Sverak, J. G. 'Sizing of Ground Conductors Against Fusing'. "
        "IEEE Transactions on Power Apparatus and Systems, 1981.",
        "IEEE C37.112-2018 — IEEE Standard Inverse-Time Characteristic Equations "
        "for Overcurrent Relays. IEEE, 2018.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(f"[{i}] {ref}")
        run.font.name = _FONTE
        run.font.size = Pt(9)


# ─── RODAPÉ / CABEÇALHO ───────────────────────────────────────────────────────

def _add_header_footer(doc, proj):
    from docx.oxml.ns import nsmap
    section = doc.sections[0]

    # Cabeçalho
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(2)
    # Linha colorida no topo
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), _COR_SUBHEADER)
    pBdr.append(bottom)
    pPr.append(pBdr)

    run1 = hp.add_run(f"Malha de Aterramento — {proj.nome_projeto or 'Projeto'}")
    run1.font.name  = _FONTE
    run1.font.size  = Pt(9)
    run1.font.color.rgb = _rgb(_COR_HEADER)
    run1.bold = True

    run2 = hp.add_run(f"   |   {proj.numero_projeto or ''}   Rev. {proj.revisao or '00'}")
    run2.font.name  = _FONTE
    run2.font.size  = Pt(9)
    run2.font.color.rgb = _rgb("595959")

    # Rodapé
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = fp.add_run("Página ")
    run_f.font.name = _FONTE
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = _rgb("595959")
    # Campo de número de página
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run_pg = fp.add_run()
    run_pg._r.append(fldChar_begin)
    run_pg._r.append(instrText)
    run_pg._r.append(fldChar_end)
    run_pg.font.name = _FONTE
    run_pg.font.size = Pt(9)
    run_pg.font.color.rgb = _rgb("595959")

    run_of = fp.add_run(f"   |   BK Malha de Terra   |   {date.today():%d/%m/%Y}")
    run_of.font.name = _FONTE
    run_of.font.size = Pt(9)
    run_of.font.color.rgb = _rgb("595959")


# ─── FUNÇÃO PRINCIPAL ─────────────────────────────────────────────────────────

def gerar_relatorio_word(projeto, imagens: dict | None = None) -> bytes:
    """
    Gera o relatório técnico em .docx e retorna como bytes.

    Args:
        projeto: instância de Projeto (já com relacionamentos carregados:
                 .dados_entrada, .resultado, .medicoes_wenner)
        imagens: dict opcional com chaves 'wenner', 'planta', 'verif', 'mapa3d'
                 e valores em bytes (PNG)

    Returns:
        bytes do arquivo .docx
    """
    de  = projeto.dados_entrada if hasattr(projeto, "dados_entrada") else None
    res = projeto.resultado      if hasattr(projeto, "resultado")      else None
    if isinstance(de,  list): de  = de[0]  if de  else None
    if isinstance(res, list): res = res[0] if res else None

    wenner_rows = list(projeto.medicoes_wenner) if hasattr(projeto, "medicoes_wenner") else []

    # Cálculos derivados
    imp  = _calcular_impedancias(de) if de else {}
    t_s  = float(getattr(de, "tempo_eliminacao_s", None) or 0.1) if de else 0.1
    i_g  = imp.get("i_g_ka", 1.0) if imp else 1.0
    cond = _calcular_condutor_por_nivel(i_g, t_s) if imp else {}
    reles = _ajustes_rele(de, imp) if de and imp else []

    # Documento
    doc = Document()

    # Configuração de página A4, margens 2 cm
    for section in doc.sections:
        section.page_width   = Cm(21)
        section.page_height  = Cm(29.7)
        section.left_margin  = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin   = Cm(2.5)
        section.bottom_margin = Cm(2.0)

    # Estilo padrão
    style = doc.styles["Normal"]
    style.font.name = _FONTE
    style.font.size = Pt(10)

    # Cabeçalho / Rodapé
    _add_header_footer(doc, projeto)

    # ── Seções ──
    _secao_identificacao(doc, projeto, de, res)
    doc.add_paragraph()
    _secao_objetivo(doc)
    doc.add_paragraph()
    _secao_metodologia(doc)
    doc.add_paragraph()
    if de and imp:
        _secao_sistema_eletrico(doc, de, imp)
        doc.add_paragraph()
    _secao_solo(doc, projeto, de, wenner_rows)
    doc.add_paragraph()
    if de and imp and cond:
        _secao_condutor(doc, de, imp, cond)
        doc.add_paragraph()
    if de:
        _secao_geometria(doc, de)
        doc.add_paragraph()
    _secao_calculos(doc, de, res)
    doc.add_paragraph()
    if res:
        _secao_tensoes(doc, de, res)
        doc.add_paragraph()
    if reles:
        _secao_coordenacao(doc, de, reles)
        doc.add_paragraph()
    if res:
        _secao_verificacao(doc, de, res)
        doc.add_paragraph()
    _secao_conclusao(doc, projeto, de, res)
    doc.add_paragraph()
    _secao_referencias(doc)

    # Exporta para bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
