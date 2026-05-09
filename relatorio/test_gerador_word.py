"""
tests/test_gerador_word.py
==========================

Testa geração do relatório Word com mock de projeto (sem banco).
"""

from datetime import date
from types import SimpleNamespace
from io import BytesIO
import pytest

from relatorio.gerador_word import gera_relatorio_word, nome_arquivo_padrao


def _mock_projeto():
    """Cria um projeto fake com a estrutura do ORM."""
    p = SimpleNamespace(
        id=1,
        cliente="Cliente Teste S.A.",
        nome_projeto="SE Industrial 138/13.8 kV",
        numero_projeto="BK-2026-001",
        revisao="00",
        responsavel_tecnico="Eng. Velho",
        crea_responsavel="PR-123456/D",
        concessionaria="Celesc",
        data_calculo=date(2026, 5, 9),
        observacoes=None,
    )
    p.medicoes_wenner = [
        SimpleNamespace(ponto=1, espacamento_m=1.0,
                          resistencia_ohm=47.7, rho_aparente=299.7),
        SimpleNamespace(ponto=2, espacamento_m=2.0,
                          resistencia_ohm=22.3, rho_aparente=280.2),
        SimpleNamespace(ponto=3, espacamento_m=4.0,
                          resistencia_ohm=9.5, rho_aparente=238.8),
        SimpleNamespace(ponto=4, espacamento_m=8.0,
                          resistencia_ohm=3.7, rho_aparente=185.9),
    ]
    p.dados_entrada = SimpleNamespace(
        largura_m=40, comprimento_m=50, profundidade_malha_m=0.5,
        espac_malha_principal_m=5, espac_malha_juncao_m=2.5,
        brita_espessura_m=0.10, brita_resistividade_ohm=3000,
        haste_comprimento_m=3, haste_diametro_mm=15.875,
        condutor_material="cobre_nu", condutor_bitola_mm2=50,
        i_falta_3i0_ka=8.0, tempo_eliminacao_s=0.5,
        sf_div_corrente=0.6, xr_ratio=12, df_decremento=1.03,
        peso_pessoa_kg=50,
    )
    p.resultado = SimpleNamespace(
        rho1_ohm_m=285.9, rho2_ohm_m=111.6, h1_m=4.57, rho_equivalente=285.9,
        bitola_calculada_mm2=20.7, bitola_adotada_mm2=50,
        cs_brita=0.719, etoque_admissivel_v=695, epasso_admissivel_v=2287,
        df_decremento=1.03, ig_corrente_malha_a=4950,
        rg_sverak_ohm=3.044, rg_schwarz_ohm=3.088, rg_adotado_ohm=3.088,
        gpr_v=15286,
        em_tensao_malha_v=600, es_tensao_passo_v=1500,
        num_hastes=24, comprimento_total_cabo_m=890,
        atende_toque=True, atende_passo=True, atende_geral=True,
        margem_toque_pct=13.7, margem_passo_pct=34.4,
    )
    p.relatorios = []
    return p


def test_gera_relatorio_basico():
    """Gera relatório sem imagens e valida que é um docx válido."""
    projeto = _mock_projeto()
    docx_bytes = gera_relatorio_word(projeto, imagens={})
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 5000  # > 5KB
    # Verifica magic number do ZIP (docx é ZIP)
    assert docx_bytes[:2] == b"PK"


def test_relatorio_abre_no_python_docx():
    """O docx gerado deve ser legível pelo próprio python-docx."""
    from docx import Document
    projeto = _mock_projeto()
    docx_bytes = gera_relatorio_word(projeto, imagens={})
    doc = Document(BytesIO(docx_bytes))
    # Deve ter pelo menos os headings principais
    todos_textos = "\n".join(p.text for p in doc.paragraphs)
    for esperado in [
        "MEMÓRIA DE CÁLCULO",
        "1. Objetivo",
        "2. Metodologia",
        "3. Dados de Entrada",
        "4. Resultados",
        "5. Conclusão",
        "6. Referências",
    ]:
        assert esperado in todos_textos, f"Faltou seção: {esperado}"


def test_relatorio_nao_atende():
    """Conclusão deve ser diferente quando malha não atende."""
    from docx import Document
    projeto = _mock_projeto()
    projeto.resultado.atende_geral = False
    projeto.resultado.atende_toque = False
    projeto.resultado.margem_toque_pct = -25.0
    docx_bytes = gera_relatorio_word(projeto, imagens={})
    doc = Document(BytesIO(docx_bytes))
    textos = "\n".join(p.text for p in doc.paragraphs)
    assert "NÃO atende" in textos


def test_nome_arquivo_padrao():
    projeto = _mock_projeto()
    nome = nome_arquivo_padrao(projeto)
    assert nome.startswith("MC_MALHA_BK")
    assert nome.endswith(".docx")
    assert "R00" in nome
    assert " " not in nome  # Sem espaços


def test_falha_sem_resultado():
    projeto = _mock_projeto()
    projeto.resultado = None
    with pytest.raises(ValueError, match="resultados"):
        gera_relatorio_word(projeto, imagens={})


def test_caracteres_especiais_no_cliente():
    """Aceita caracteres especiais e gera nome de arquivo seguro."""
    projeto = _mock_projeto()
    projeto.cliente = "Cliente / Especial & Cia Ltda."
    docx_bytes = gera_relatorio_word(projeto, imagens={})
    assert len(docx_bytes) > 5000

    nome = nome_arquivo_padrao(projeto)
    assert "/" not in nome
    assert "&" not in nome
    assert "." not in nome.replace(".docx", "")


def test_estrutura_ooxml_valida():
    """
    Valida que o XML do .docx está conforme schema OOXML:
    - Filhos diretos do <w:body> só podem ser <w:p>, <w:tbl> ou <w:sectPr>
    - <w:sectPr> deve ser o ÚLTIMO filho do body
    - Equações OMML (<m:oMathPara>) devem estar DENTRO de <w:p>

    Esse teste protege contra o bug de inserir oMathPara solto no body,
    que faz o Word recusar o arquivo com erro genérico de "permissão".
    """
    import zipfile
    from io import BytesIO

    try:
        from lxml import etree
    except ImportError:
        pytest.skip("lxml não instalado")

    projeto = _mock_projeto()
    docx_bytes = gera_relatorio_word(projeto, imagens={})

    # Parse XML
    with zipfile.ZipFile(BytesIO(docx_bytes)) as z:
        with z.open('word/document.xml') as f:
            tree = etree.parse(f)

    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
    }

    # 1. Filhos do body só podem ser p, tbl ou sectPr
    body = tree.xpath('//w:body', namespaces=ns)[0]
    tags_permitidas = {
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p',
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl',
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr',
    }
    for child in body:
        assert child.tag in tags_permitidas, (
            f"Elemento inválido como filho direto de <w:body>: {child.tag}"
        )

    # 2. sectPr deve ser o último (Word exige isso)
    ultimo_tag = body[-1].tag
    assert ultimo_tag.endswith('}sectPr'), (
        f"Último filho do body deve ser <w:sectPr>, encontrado: {ultimo_tag}"
    )

    # 3. Todo oMathPara deve estar dentro de um <w:p>
    omath_paras = tree.xpath('//m:oMathPara', namespaces=ns)
    for op in omath_paras:
        parent = op.getparent()
        assert parent.tag.endswith('}p'), (
            f"<m:oMathPara> com parent inválido: {parent.tag} (deveria ser <w:p>)"
        )
